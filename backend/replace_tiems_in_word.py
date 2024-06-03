from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import qrcode
from PIL import Image

def create_qr_code(url, qr_path):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    img.save(qr_path)

def replace_text_in_docx(input_path, output_path, replacements, qr_url):
    # Открываем исходный документ
    doc = Document(input_path)

    # Проходим по всем параграфам документа
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements:
            if old_text in paragraph.text:
                new_paragraph_text = paragraph.text.replace(old_text, new_text)
                paragraph.clear()

                # Создаем новый пробег для всего текста
                run = paragraph.add_run(new_paragraph_text)

                # Проходим по каждому замененному тексту и настраиваем его шрифт и размер
                for old_text, new_text in replacements:
                    if new_text in run.text:
                        start_index = run.text.find(new_text)
                        end_index = start_index + len(new_text)

                        # Разделяем текст на части
                        run_before = run.text[:start_index]
                        run_middle = new_text
                        run_after = run.text[end_index:]

                        # Удаляем исходный пробег
                        paragraph.clear()

                        # Добавляем новый пробег до замененного текста
                        if run_before:
                            paragraph.add_run(run_before)

                        # Добавляем новый пробег с замененным текстом
                        middle_run = paragraph.add_run(run_middle)
                        middle_run.font.name = 'Arial'
                        middle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        middle_run.font.size = Pt(10.5)
                        if old_text in ['NAME', 'SUMMA']:
                            middle_run.bold = True

                        # Добавляем новый пробег после замененного текста
                        if run_after:
                            paragraph.add_run(run_after)

    # Создаем и добавляем QR-код в конец документа
    qr_path = 'qr_code.png'
    create_qr_code(qr_url, qr_path)
    doc.add_paragraph('QR Code:')
    doc.add_picture(qr_path, width=Inches(2.0))

    # Сохраняем измененный документ
    doc.save(output_path)

def replace_text_in_docx_ariza(input_path, output_path, replacements):
    # Открываем исходный документ
    doc = Document(input_path)

    # Проходим по всем параграфам документа
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements:
            if old_text in paragraph.text:
                new_paragraph_text = paragraph.text.replace(old_text, new_text)
                paragraph.clear()

                # Создаем новый пробег для всего текста
                run = paragraph.add_run(new_paragraph_text)

                # Проходим по каждому замененному тексту и настраиваем его шрифт и размер
                for old_text, new_text in replacements:
                    if new_text in run.text:
                        start_index = run.text.find(new_text)
                        end_index = start_index + len(new_text)

                        # Разделяем текст на части
                        run_before = run.text[:start_index]
                        run_middle = new_text
                        run_after = run.text[end_index:]

                        # Удаляем исходный пробег
                        paragraph.clear()

                        # Добавляем новый пробег до замененного текста
                        if run_before:
                            paragraph.add_run(run_before)

                        # Добавляем новый пробег с замененным текстом
                        middle_run = paragraph.add_run(run_middle)
                        middle_run.font.name = 'Arial'
                        middle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        middle_run.font.size = Pt(10.5)
                        if old_text in ['NAME']:
                            middle_run.bold = True

                        # Добавляем новый пробег после замененного текста
                        if run_after:
                            paragraph.add_run(run_after)

    # Сохраняем измененный документ
    doc.save(output_path)